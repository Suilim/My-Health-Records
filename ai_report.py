from datetime import date, timedelta
from collections import defaultdict
from io import BytesIO

import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from export_records import get_user_records
from settings_utils import get_drug_slots, get_recorded_slots_by_date


# ── 工具函式 ──────────────────────────────────────────────

def _week_label(d: date) -> str:
    week_num = (d.day - 1) // 7 + 1
    return f"{d.month}月第{week_num}週"


def _week_range_label(week_start: date, week_end: date) -> str:
    return f"{week_start.month}/{week_start.day}～{week_end.month}/{week_end.day}"


def _iter_weeks(start_date: date, end_date: date):
    """逐週產生 (week_start, week_end)，週一為起點"""
    # 對齊到週一
    current = start_date - timedelta(days=start_date.weekday())
    while current <= end_date:
        week_end = min(current + timedelta(days=6), end_date)
        week_start = max(current, start_date)
        yield week_start, week_end
        current += timedelta(days=7)


def _dates_in_range(week_start: date, week_end: date):
    d = week_start
    while d <= week_end:
        yield d
        d += timedelta(days=1)


# ── 各模組的週塊資料整理 ──────────────────────────────────

def _block_drug(user_id: str, week_start: date, week_end: date,
                all_drug_records: list) -> list[str]:
    """
    回傳該週用藥的行列表。
    - required_slots：使用者設定的時段（早/午/晚/睡前）
    - 全週正常 → 一行「用藥：全週正常 ✅」
    - 有漏 → 列出漏服的日期+時段
    - 需要時 → 單獨列出
    """
    required_slots = get_drug_slots(user_id)
    recorded_slots = get_recorded_slots_by_date(user_id)  # {"2026-01-05": {"早", "晚"}, ...}

    lines = []

    # ── 正規時段漏填檢查 ──
    if required_slots:
        missing_entries = []
        for d in _dates_in_range(week_start, week_end):
            d_str = d.strftime("%Y-%m-%d")
            filled = recorded_slots.get(d_str, set())
            for slot in required_slots:
                if slot not in filled:
                    missing_entries.append(f"{d.month}/{d.day} {slot}❌")

        if missing_entries:
            lines.append(f"用藥：{'、'.join(missing_entries)}")
        else:
            lines.append("用藥：全週正常 ✅")

    # ── 需要時 ──
    prn_drugs = []
    for r in all_drug_records:
        try:
            d_str = r["filltime"][:10]
            d = date.fromisoformat(d_str)
            if week_start <= d <= week_end and r.get("eattime", "").strip() == "需要時":
                name = r.get("drugname", "").strip()
                pieces = r.get("drugpieces", "")
                label = f"{d.month}/{d.day} {name}"
                if pieces:
                    label += f" {pieces}顆"
                prn_drugs.append(label)
        except Exception:
            continue

    if prn_drugs:
        lines.append(f"需要時：{'、'.join(prn_drugs)}")

    return lines


def _block_symptom(week_start: date, week_end: date, all_symptom_records: list) -> list[str]:
    """不舒服：症狀名稱（次數，好發時間）"""
    symptom_data = defaultdict(lambda: {"count": 0, "times": []})

    for r in all_symptom_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            name = r.get("symptomname", "").strip()
            if not name or name == "（無症狀）":
                continue
            symptom_data[name]["count"] += 1
            occur_time = r.get("occurtime", "").strip()
            if occur_time:
                symptom_data[name]["times"].append(occur_time)
        except Exception:
            continue

    if not symptom_data:
        return []

    parts = []
    for name, info in symptom_data.items():
        count = info["count"]
        times = info["times"]
        if times:
            # 取最常見的發生時間
            most_common = max(set(times), key=times.count)
            parts.append(f"{name}（{count}次，{most_common}）")
        else:
            parts.append(f"{name}（{count}次）")

    return [f"不舒服：{'、'.join(parts)}"]


def _block_weight(week_start: date, week_end: date, all_weight_records: list) -> list[str]:
    """體重：週初→週末"""
    week_weights = []
    for r in all_weight_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            w = float(r.get("wei", 0))
            if w > 0:
                week_weights.append((d, w))
        except Exception:
            continue

    if not week_weights:
        return []

    week_weights.sort()
    if len(week_weights) == 1:
        return [f"體重：{week_weights[0][1]} kg"]

    start_w = week_weights[0][1]
    end_w = week_weights[-1][1]
    return [f"體重：{start_w} → {end_w} kg"]


def _block_life(week_start: date, week_end: date, all_life_records: list) -> list[str]:
    """生活：情緒統計 + 日記全文合併"""
    emotions = []
    diaries = []

    for r in all_life_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            emotion = r.get("emotion", "").strip()
            if emotion:
                for e in emotion.split("、"):
                    e = e.strip()
                    if e:
                        emotions.append(e)
            diary = r.get("liferecord", "").strip()
            if diary:
                diaries.append(diary)
        except Exception:
            continue

    if not emotions and not diaries:
        return []

    parts = []
    if emotions:
        emotion_count = defaultdict(int)
        for e in emotions:
            emotion_count[e] += 1
        emotion_str = "、".join(f"{e}×{n}" for e, n in emotion_count.items())
        parts.append(f"情緒〔{emotion_str}〕")

    if diaries:
        diary_text = " ".join(diaries)
        parts.append(f"日記：{diary_text}")

    return [f"生活：{'　'.join(parts)}"]


def _block_sleep(week_start: date, week_end: date, all_sleep_records: list) -> list[str]:
    """睡眠：週平均時數與品質"""
    durations = []
    qualities = []

    for r in all_sleep_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            dur = r.get("duration")
            if dur is not None:
                durations.append(float(dur))
            qual = r.get("quality")
            if qual is not None:
                qualities.append(float(qual))
        except Exception:
            continue

    if not durations and not qualities:
        return []

    parts = []
    if durations:
        parts.append(f"平均 {sum(durations)/len(durations):.1f} 小時")
    if qualities:
        parts.append(f"品質 {sum(qualities)/len(qualities):.1f}/5")

    return [f"睡眠：{'，'.join(parts)}"]


def _block_simple(week_start: date, week_end: date, all_records: list,
                   field: str, display_name: str, unit: str,
                   high_threshold: float = None, low_threshold: float = None) -> list[str]:
    """血壓/血糖/體溫：週平均 + 標異常日期"""
    daily_vals = defaultdict(list)

    for r in all_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            val = float(r.get(field, 0))
            if val > 0:
                daily_vals[d].append(val)
        except Exception:
            continue

    if not daily_vals:
        return []

    all_vals = [v for vals in daily_vals.values() for v in vals]
    avg = sum(all_vals) / len(all_vals)
    line = f"{display_name}：週平均 {avg:.1f} {unit}"

    # 標注異常日期
    abnormal = []
    for d, vals in sorted(daily_vals.items()):
        day_avg = sum(vals) / len(vals)
        if high_threshold and day_avg > high_threshold:
            abnormal.append(f"{d.month}/{d.day} {day_avg:.0f}偏高⚠️")
        elif low_threshold and day_avg < low_threshold:
            abnormal.append(f"{d.month}/{d.day} {day_avg:.0f}偏低⚠️")

    if abnormal:
        line += f"　{'、'.join(abnormal)}"

    return [line]


def _block_heartrate(week_start: date, week_end: date, all_records: list) -> list[str]:
    """血壓：週平均收縮壓/舒張壓 + 標異常日期"""
    daily_sys = defaultdict(list)
    daily_dia = defaultdict(list)

    for r in all_records:
        try:
            d = date.fromisoformat(r["filltime"][:10])
            if not (week_start <= d <= week_end):
                continue
            sys = float(r.get("mmHg1", 0))
            dia = float(r.get("mmHg2", 0))
            if sys > 0:
                daily_sys[d].append(sys)
                daily_dia[d].append(dia)
        except Exception:
            continue

    if not daily_sys:
        return []

    all_sys = [v for vals in daily_sys.values() for v in vals]
    all_dia = [v for vals in daily_dia.values() for v in vals]
    avg_sys = sum(all_sys) / len(all_sys)
    avg_dia = sum(all_dia) / len(all_dia)
    line = f"血壓：週平均 {avg_sys:.0f}/{avg_dia:.0f} mmHg"

    abnormal = []
    for d in sorted(daily_sys.keys()):
        day_sys = sum(daily_sys[d]) / len(daily_sys[d])
        if day_sys > 130:
            abnormal.append(f"{d.month}/{d.day} {day_sys:.0f}偏高⚠️")
        elif day_sys < 90:
            abnormal.append(f"{d.month}/{d.day} {day_sys:.0f}偏低⚠️")

    if abnormal:
        line += f"　{'、'.join(abnormal)}"

    return [line]


# ── 主要整理函式 ──────────────────────────────────────────

def prepare_data_for_ai(user_id: str, start_date: date, end_date: date,
                         selected_modules: list) -> str:
    """
    從 Firebase 讀取各模組資料，整理成以週為單位的純文字給 Gemini。
    selected_modules: 模組 key 列表，如 ['heartrate', 'drug', 'life', ...]
    """
    MODULE_TO_NODE = {
        "heartrate": "HeartRate",
        "weight": "Weight",
        "sugar": "Sugar",
        "temp": "Temp",
        "drug": "Drug",
        "life": "Life",
        "symptom": "Symptom",
        "sleep": "Sleep",
    }

    # 一次取出所有模組的原始資料
    all_records = {}
    for mod in selected_modules:
        node = MODULE_TO_NODE.get(mod)
        if node:
            all_records[mod] = get_user_records(user_id, node, start_date, end_date)

    # 逐週建立區塊
    week_blocks = []
    for week_start, week_end in _iter_weeks(start_date, end_date):
        label = f"### {_week_label(week_start)}（{_week_range_label(week_start, week_end)}）"
        lines = [label]

        if "drug" in selected_modules:
            lines += _block_drug(user_id, week_start, week_end, all_records.get("drug", []))

        if "symptom" in selected_modules:
            lines += _block_symptom(week_start, week_end, all_records.get("symptom", []))

        if "weight" in selected_modules:
            lines += _block_weight(week_start, week_end, all_records.get("weight", []))

        if "sleep" in selected_modules:
            lines += _block_sleep(week_start, week_end, all_records.get("sleep", []))

        if "life" in selected_modules:
            lines += _block_life(week_start, week_end, all_records.get("life", []))

        if "heartrate" in selected_modules:
            lines += _block_heartrate(week_start, week_end, all_records.get("heartrate", []))

        if "sugar" in selected_modules:
            lines += _block_simple(week_start, week_end, all_records.get("sugar", []),
                                   "sugarlevel", "血糖", "mg/dL",
                                   high_threshold=126, low_threshold=70)

        if "temp" in selected_modules:
            lines += _block_simple(week_start, week_end, all_records.get("temp", []),
                                   "temp", "體溫", "°C",
                                   high_threshold=37.5, low_threshold=36.0)

        # 如果該週除了標題外沒有任何資料就略過
        if len(lines) > 1:
            week_blocks.append("\n".join(lines))

    if not week_blocks:
        return "（此期間無任何紀錄）"

    return "\n\n".join(week_blocks)


# ── Gemini 呼叫 ───────────────────────────────────────────

def generate_report_with_gemini(user_name: str, data_context: str,
                                 start_date: date, end_date: date) -> str:
    """呼叫 Gemini API，回傳 markdown 格式的健康報告。"""
    api_key = st.secrets.get("gemini", {}).get("api_key", "")
    if not api_key:
        raise ValueError("請在 .streamlit/secrets.toml 設定 [gemini] api_key")

    client = genai.Client(api_key=api_key)

    prompt = f"""你是一位專業醫療助理，請根據以下健康紀錄，撰寫一份給醫師閱覽的健康狀況報告。

使用者：{user_name}
紀錄期間：{start_date.strftime("%Y年%m月%d日")} 至 {end_date.strftime("%Y年%m月%d日")}

健康紀錄資料（以週為單位）：
---
{data_context}
---

請依以下格式撰寫報告（繁體中文，語氣專業但易讀）：

# {user_name} 的狀況記錄
**紀錄期間：{start_date.strftime("%Y/%m/%d")} ～ {end_date.strftime("%Y/%m/%d")}**

## 各週摘要
（逐週說明重要事項，略過平淡無事的週次）

## 跨週關聯分析
（找出跨週的關聯，例如：某週日記提到月經來潮，同週頭痛頻繁；某段時間血壓偏高與睡眠品質下降同時出現等。重點放在可能對醫師有參考價值的觀察）

## 整體觀察
（2-4句整體觀察，供醫師參考，非診斷）
"""

    response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents=prompt,
    )
    return response.text


# ── Word 匯出 ─────────────────────────────────────────────

def export_report_to_docx(report_text: str, user_name: str,
                           start_date: date, end_date: date) -> BytesIO:
    """將 markdown 格式報告轉成 Word 檔，回傳 BytesIO。"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微軟正黑體"
    style.font.size = Pt(11)

    for line in report_text.split("\n"):
        line = line.rstrip()

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "":
            doc.add_paragraph("")
        else:
            if "**" in line:
                p = doc.add_paragraph()
                parts = line.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
